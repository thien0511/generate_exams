import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
import io
import random
import base64
import docx
from numpy import int64

# Set page title and configuration
st.set_page_config(
    page_title="Ph·∫ßn M·ªÅm T·∫°o ƒê·ªÅ Thi",
    page_icon="üìù",
    layout="wide"
)

# Title and description
st.title("Ph·∫ßn M·ªÅm T·∫°o ƒê·ªÅ Thi")
st.markdown("""
·ª®ng d·ª•ng n√†y cho ph√©p b·∫°n:
1. T·∫£i l√™n th∆∞ vi·ªán c√¢u h·ªèi Excel c·ªßa b·∫°n
2a. Xu·∫•t t·∫•t c·∫£ c√¢u h·ªèi ra Word
2b. Ch·ªçn ng·∫´u nhi√™n m·ªôt s·ªë c√¢u h·ªèi
3. X√°o tr·ªôn c√¢u h·ªèi v√† ƒë√°p √°n  
4. Xu·∫•t nh·ªØng c√¢u h·ªèi ƒë√≥ ra t√†i li·ªáu MS Word.
""")

# Initialize session state variables
if 'question_library' not in st.session_state:
    st.session_state.question_library = None
if 'file_uploaded' not in st.session_state:
    st.session_state.file_uploaded = False
if 'file_valid' not in st.session_state:
    st.session_state.file_valid = False
if 'selected_questions' not in st.session_state:
    st.session_state.selected_questions = None
if 'randomized_questions' not in st.session_state:
    st.session_state.randomized_questions = None
if 'docx_data' not in st.session_state:
    st.session_state.docx_data = None
if 'trigger_download' not in st.session_state:
    st.session_state.trigger_download = False

# File upload section
st.header("T·∫£i L√™n Th∆∞ Vi·ªán C√¢u H·ªèi")
st.markdown("T·∫£i l√™n m·ªôt t·ªáp Excel (.xlsx ho·∫∑c .xls) ch·ª©a th∆∞ vi·ªán c√¢u h·ªèi c·ªßa b·∫°n. T·ªáp ph·∫£i c√≥ √≠t nh·∫•t 10 c·ªôt. C√°c c·ªôt theo th·ª© t·ª± l√†: 'Ma cau', 'Cau hoi', 'Tra loi 1', 'Tra loi 2', 'Tra loi 3', 'Tra loi 4', 'Dap an dung', 'Bai', 'Phan', 'Do kho'.")

uploaded_file = st.file_uploader("Ch·ªçn t·ªáp Excel", type=['xlsx', 'xls'])

if uploaded_file is not None:
    try:
        # Attempt to read the Excel file
        df = pd.read_excel(uploaded_file)
        
        # Define the column names
        column_names = [
            "Ma cau", "Cau hoi", "Tra loi 1", "Tra loi 2", "Tra loi 3", 
            "Tra loi 4", "Dap an dung", "Bai", "Phan", "Do kho"
        ]
        
        # Check if the file has at least 10 columns
        if len(df.columns) >= 10:
            # Select only the first 10 columns and rename them
            df = df.iloc[:, :10]
            df.columns = column_names
            st.success("T·ªáp ƒë√£ ƒë∆∞·ª£c t·∫£i l√™n v√† ki·ªÉm tra th√†nh c√¥ng!")
            st.session_state.question_library = df
            st.session_state.file_uploaded = True
            st.session_state.file_valid = True
            
            # Display a preview of the uploaded data
            st.subheader("Xem tr∆∞·ªõc th∆∞ vi·ªán c√¢u h·ªèi")
            st.dataframe(df)
            
            # Display total number of questions in the library
            st.info(f"T·ªïng s·ªë c√¢u h·ªèi trong th∆∞ vi·ªán: {len(df)}")
            
            # Add button to export all questions
            if st.button("**Xu·∫•t t·∫•t c·∫£ c√¢u h·ªèi ra Word**"):
                # Create Word document
                doc = Document()
                
                # Apply single line spacing with no spacing before/after
                style = doc.styles['Normal']
                style.paragraph_format.line_spacing = 1.0  # Single spacing
                style.paragraph_format.space_before = 0
                style.paragraph_format.space_after = 0
                
                # Add page numbers to the footer
                section = doc.sections[0]
                footer = section.footer
                paragraph = footer.paragraphs[0]
                paragraph.alignment = 2  # Right alignment
                run = paragraph.add_run()
                fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
                fldChar1.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
                run._element.append(fldChar1)
                
                instrText = docx.oxml.shared.OxmlElement('w:instrText')
                instrText.set(docx.oxml.shared.qn('xml:space'), 'preserve')
                instrText.text = " PAGE "
                run._element.append(instrText)
                
                fldChar2 = docx.oxml.shared.OxmlElement('w:fldChar')
                fldChar2.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
                run._element.append(fldChar2)
                
                # Store answer key information for adding at the end
                answer_key = []
                
                # Add each question to the document
                for i, row in enumerate(df.itertuples(), 1):
                    # Create a single paragraph for question number and text
                    question_para = doc.add_paragraph()
                    question_para.paragraph_format.line_spacing = 1.0
                    # Add space before if not the first question
                    if i > 1:
                        question_para.paragraph_format.space_before = Pt(0)
                    else:
                        question_para.paragraph_format.space_before = 0
                    question_para.paragraph_format.space_after = 0
                    
                    # Add the question number in bold
                    question_number = question_para.add_run(f"C√¢u {i}: ")
                    question_number.bold = True
                    
                    # Add the question text in italic
                    question_text = question_para.add_run(row.Cau_hoi if hasattr(row, 'Cau_hoi') else row[2])
                    question_text.italic = True
                    
                    # Get answer options and correct answer
                    options = [
                        ("A", row.Tra_loi_1 if hasattr(row, 'Tra_loi_1') else row[3]),
                        ("B", row.Tra_loi_2 if hasattr(row, 'Tra_loi_2') else row[4]),
                        ("C", row.Tra_loi_3 if hasattr(row, 'Tra_loi_3') else row[5]),
                        ("D", row.Tra_loi_4 if hasattr(row, 'Tra_loi_4') else row[6])
                    ]
                    correct_answer = row.Dap_an_dung if hasattr(row, 'Dap_an_dung') else row[7]
                    
                    # Convert numeric correct answer to letter if needed
                    if isinstance(correct_answer, (int, float)):
                        try:
                            numeric_answer = int(correct_answer)
                            if 1 <= numeric_answer <= 4:
                                correct_answer = chr(64 + numeric_answer)  # Convert 1->A, 2->B, 3->C, 4->D
                        except (ValueError, TypeError):
                            pass
                    
                    # Store answer key information for adding at the end
                    answer_key.append(f"{i}{correct_answer}")
                    
                    # Calculate the layout based on answer lengths
                    answer_lengths = [len(str(content)) for content in [opt[1] for opt in options]]
                    avg_length = sum(answer_lengths) / len(answer_lengths)
                    max_length = max(answer_lengths)
                    
                    # Decide layout based on answer lengths
                    if max_length < 15 and avg_length < 10:  # Short answers - put all on one line
                        answers_para = doc.add_paragraph()
                        answers_para.paragraph_format.line_spacing = 1.0
                        answers_para.paragraph_format.space_before = 0
                        answers_para.paragraph_format.space_after = 0
                        
                        # Calculate spacing
                        page_width_points = 450
                        total_answer_text_len = sum(len(str(option_text)) for _, option_text in options)
                        total_letter_len = 8
                        space_to_distribute = max(20, page_width_points - total_answer_text_len - total_letter_len)
                        space_between = min(20, int(space_to_distribute / 3))
                        spacing = " " * space_between
                        
                        for idx, (option_letter, option_text) in enumerate(options):
                            if idx > 0:
                                answers_para.add_run(spacing)
                            
                            option_label = answers_para.add_run(f"{option_letter}. ")
                            option_label.bold = True
                            
                            is_correct = (option_letter == correct_answer)
                            option_content = answers_para.add_run(str(option_text))
                            if is_correct:
                                option_content.underline = True
                                
                    elif max_length < 40 and avg_length < 30:  # Medium length - 2x2 layout
                        # First row (A and B)
                        row1_para = doc.add_paragraph()
                        row1_para.paragraph_format.line_spacing = 1.0
                        row1_para.paragraph_format.space_before = 0
                        row1_para.paragraph_format.space_after = 0
                        
                        page_width_points = 450
                        row1_text_len = sum(len(str(option_text)) for _, option_text in options[:2])
                        letter_len = 4
                        space_to_distribute = max(20, page_width_points - row1_text_len - letter_len)
                        space_between = min(20, int(space_to_distribute))
                        spacing = " " * space_between
                        
                        for idx in range(2):
                            if idx > 0:
                                row1_para.add_run(spacing)
                            
                            option_letter, option_text = options[idx]
                            option_label = row1_para.add_run(f"{option_letter}. ")
                            option_label.bold = True
                            
                            is_correct = (option_letter == correct_answer)
                            option_content = row1_para.add_run(str(option_text))
                            if is_correct:
                                option_content.underline = True
                        
                        # Second row (C and D)
                        row2_para = doc.add_paragraph()
                        row2_para.paragraph_format.line_spacing = 1.0
                        row2_para.paragraph_format.space_before = 0
                        row2_para.paragraph_format.space_after = 0
                        
                        row2_text_len = sum(len(str(option_text)) for _, option_text in options[2:4])
                        space_to_distribute = max(20, page_width_points - row2_text_len - letter_len)
                        space_between = min(20, int(space_to_distribute))
                        spacing = " " * space_between
                        
                        for idx in range(2, 4):
                            if idx > 2:
                                row2_para.add_run(spacing)
                            
                            option_letter, option_text = options[idx]
                            option_label = row2_para.add_run(f"{option_letter}. ")
                            option_label.bold = True
                            
                            is_correct = (option_letter == correct_answer)
                            option_content = row2_para.add_run(str(option_text))
                            if is_correct:
                                option_content.underline = True
                    
                    else:  # Long answers - one per line
                        for option_letter, option_text in options:
                            answer_para = doc.add_paragraph()
                            answer_para.paragraph_format.line_spacing = 1.0
                            answer_para.paragraph_format.space_before = 0
                            answer_para.paragraph_format.space_after = 0
                            
                            option_label = answer_para.add_run(f"{option_letter}. ")
                            option_label.bold = True
                            
                            is_correct = (option_letter == correct_answer)
                            option_content = answer_para.add_run(str(option_text))
                            if is_correct:
                                option_content.underline = True
                
                # Add the answer key at the end of the document
                if answer_key:
                    doc.add_page_break()
                    doc.add_heading('ƒê√°p √°n', level=1)
                    
                    answer_key_para = doc.add_paragraph()
                    answer_key_para.paragraph_format.line_spacing = 1.0
                    answer_key_para.paragraph_format.space_before = 0
                    answer_key_para.paragraph_format.space_after = 0
                    
                    # Format answers with proper spacing
                    formatted_answers = []
                    for answer in answer_key:
                        formatted_answers.append(answer)
                    
                    # Join answers with spaces and add to paragraph
                    answer_text = " ".join(formatted_answers)
                    answer_key_para.add_run(answer_text)
                
                # Save the document to a BytesIO object
                docx_io = io.BytesIO()
                doc.save(docx_io)
                docx_io.seek(0)
                
                # Convert the document to base64
                docx_base64 = base64.b64encode(docx_io.getvalue()).decode()
                
                # Create JavaScript to trigger download
                js = f"""
                <script>
                    function downloadBase64File(base64Data, fileName) {{
                        const linkSource = `data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}`;
                        const downloadLink = document.createElement("a");
                        downloadLink.href = linkSource;
                        downloadLink.download = 'all_questions.docx';
                        downloadLink.click();
                    }}
                    downloadBase64File('{docx_base64}', 'all_questions');
                </script>
                """
                
                st.success("T√†i li·ªáu Word ƒë√£ ƒë∆∞·ª£c t·∫°o th√†nh c√¥ng!")
                st.components.v1.html(js, height=0)
        else:
            st.error(f"T·ªáp ƒë√£ t·∫£i l√™n c√≥ {len(df.columns)} c·ªôt. N√≥ ph·∫£i c√≥ √≠t nh·∫•t 10 c·ªôt.")
            st.session_state.file_valid = False
    except Exception as e:
        if "Length mismatch" in str(e):
            st.error("L·ªói: S·ªë l∆∞·ª£ng c·ªôt kh√¥ng kh·ªõp. T·ªáp Excel c·ªßa b·∫°n c√≥ qu√° nhi·ªÅu c·ªôt. Vui l√≤ng ƒë·∫£m b·∫£o t·ªáp c√≥ √≠t nh·∫•t 10 c·ªôt theo th·ª© t·ª±: Ma cau, Cau hoi, Tra loi 1, Tra loi 2, Tra loi 3, Tra loi 4, Dap an dung, Bai, Phan, Do kho.")
        else:
            st.error(f"L·ªói khi ƒë·ªçc t·ªáp: {str(e)}")
        st.session_state.file_valid = False

# Question Selection and Export section (only show if file is valid)
if st.session_state.file_valid:
    st.header("T·∫°o ƒë·ªÅ thi")
    
    # Step 1: Select random questions
    st.subheader("B∆∞·ªõc 1: Ch·ªçn c√¢u h·ªèi ng·∫´u nhi√™n")
    num_questions = st.number_input(
        "S·ªë l∆∞·ª£ng c√¢u h·ªèi c·∫ßn ch·ªçn",
        min_value=1,
        max_value=len(st.session_state.question_library),
        value=min(40, len(st.session_state.question_library)),
        step=1
    )
    
    if st.button("**Ch·ªçn c√¢u h·ªèi ng·∫´u nhi√™n**"):
        if num_questions > len(st.session_state.question_library):
            st.warning(f"B·∫°n ƒë√£ y√™u c·∫ßu {num_questions} c√¢u h·ªèi, nh∆∞ng th∆∞ vi·ªán ch·ªâ c√≥ {len(st.session_state.question_library)} c√¢u h·ªèi. T·∫•t c·∫£ c√¢u h·ªèi s·∫Ω ƒë∆∞·ª£c ch·ªçn.")
            selected_questions = st.session_state.question_library
        else:
            # Randomly select questions
            selected_indices = np.random.choice(
                len(st.session_state.question_library),
                size=num_questions,
                replace=False
            )
            # Sort indices in ascending order
            selected_indices = np.sort(selected_indices)
            selected_questions = st.session_state.question_library.iloc[selected_indices].reset_index(drop=True)
        
        # Store the selected questions in session state
        st.session_state.selected_questions = selected_questions
        st.session_state.randomized_questions = None  # Reset randomized questions
        st.session_state.docx_ready = None  # Reset docx ready state
        
        # Display the selected questions
        st.subheader("C√°c c√¢u h·ªèi ƒë√£ ch·ªçn")
        st.dataframe(selected_questions)
    
    # Step 2: Randomize questions and answers
    if st.session_state.selected_questions is not None:
        st.subheader("B∆∞·ªõc 2: X√°o tr·ªôn c√¢u h·ªèi v√† ƒë√°p √°n")
        
        # Always show selected questions
        st.write("C√°c c√¢u h·ªèi ƒë√£ ch·ªçn:")
        st.dataframe(st.session_state.selected_questions)
        
        if st.button("**X√°o tr·ªôn c√¢u h·ªèi v√† ƒë√°p √°n**"):
            # Create a copy of selected questions
            randomized_questions = st.session_state.selected_questions.copy()
            
            # Randomize the order of questions
            randomized_questions = randomized_questions.sample(frac=1).reset_index(drop=True)
            
            # Randomize answers for each question
            for idx in range(len(randomized_questions)):
                # Get the original answers and correct answer
                original_answers = [
                    randomized_questions.iloc[idx]['Tra loi 1'],
                    randomized_questions.iloc[idx]['Tra loi 2'],
                    randomized_questions.iloc[idx]['Tra loi 3'],
                    randomized_questions.iloc[idx]['Tra loi 4']
                ]
                correct_answer = randomized_questions.iloc[idx]['Dap an dung']
                
                # Determine the original correct answer index
                correct_index = None
                if isinstance(correct_answer, str) and correct_answer.upper() in ["A", "B", "C", "D"]:
                    correct_index = ord(correct_answer.upper()) - ord('A')
                elif isinstance(correct_answer, (int, float, np.int64)):
                    try:
                        temp_index = int(correct_answer) - 1
                        if 0 <= temp_index <= 3:
                            correct_index = temp_index
                    except (ValueError, TypeError): 
                        pass
                
                if correct_index is None:
                    correct_index = 0
                
                # Get the correct answer content
                correct_content = original_answers[correct_index]
                
                # Shuffle the answers
                shuffled_answers = original_answers.copy()
                random.shuffle(shuffled_answers)
                
                # Find where the correct answer ended up after shuffling
                new_correct_index = shuffled_answers.index(correct_content)
                
                # Update the question with shuffled answers and new correct answer
                randomized_questions.iloc[idx, 2:6] = shuffled_answers  # Update Tra loi 1-4
                randomized_questions.iloc[idx, 6] = new_correct_index + 1 # Update Dap an dung
            
            # Store randomized questions
            st.session_state.randomized_questions = randomized_questions
        
        # Always show randomized questions if they exist
        if st.session_state.randomized_questions is not None:
            st.write("C√°c c√¢u h·ªèi ƒë√£ x√°o tr·ªôn:")
            st.dataframe(st.session_state.randomized_questions)
    
    # Step 3: Export to Word
    if st.session_state.randomized_questions is not None:
        st.subheader("B∆∞·ªõc 3: Xu·∫•t ra t√†i li·ªáu Word")
        
        # Always show selected questions
        # st.write("C√°c c√¢u h·ªèi ƒë√£ ch·ªçn:")
        # st.dataframe(st.session_state.selected_questions)
        
        # Always show randomized questions
        # st.write("C√°c c√¢u h·ªèi ƒë√£ x√°o tr·ªôn:")
        # st.dataframe(st.session_state.randomized_questions)
        
        # Option to underline correct answers
        underline_correct = st.checkbox("G·∫°ch ch√¢n ƒë√°p √°n ƒë√∫ng", value=True)
        
        # Input for space between questions
        space_between_questions = st.number_input(
            "Kho·∫£ng c√°ch gi·ªØa c√°c c√¢u (ƒëi·ªÉm)", 
            min_value=0, 
            max_value=30,
            value=0,
            help="ƒê·∫∑t kho·∫£ng c√°ch gi·ªØa c√°c c√¢u h·ªèi theo ƒëi·ªÉm (72 ƒëi·ªÉm = 1 inch). Gi√° tr·ªã cao h∆°n s·∫Ω t·∫°o ra kho·∫£ng c√°ch l·ªõn h∆°n gi·ªØa c√°c c√¢u h·ªèi."
        )
        
        # Input for file name
      
        # Add inputs for school name and exam name
        col1, col2 = st.columns(2)
        with col1:
            school_name = st.text_input(
                "T√™n Tr∆∞·ªùng",
                value="",
                help="Nh·∫≠p t√™n tr∆∞·ªùng"
            )
        with col2:
            exam_name = st.text_input(
                "T√™n b√†i ki·ªÉm tra",
                value="",
                help="Nh·∫≠p t√™n b√†i ki·ªÉm tra"
            )
        file_name = st.text_input(
            "T√™n t·ªáp Word",
            value="examination_questions",
            help="Nh·∫≠p t√™n t·ªáp Word (kh√¥ng c·∫ßn th√™m ph·∫ßn m·ªü r·ªông .docx)"
        )
          
        if st.button("**T·∫°o v√† t·∫£i xu·ªëng t√†i li·ªáu Word**"):
            # Create Word document
            doc = Document()
            
            # Add school name and exam name in a centered table only if both are provided
            if school_name.strip() and exam_name.strip():
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Normal Table'
                table.alignment = 1  # Center alignment
                
                # Set column widths
                for cell in table.columns[0].cells:
                    cell.width = Inches(2)
                for cell in table.columns[1].cells:
                    cell.width = Inches(4)
                
                # Add content to cells
                cell1 = table.cell(0, 0)
                cell2 = table.cell(0, 1)
                
                # Center align text in cells
                cell1.paragraphs[0].alignment = 1
                cell2.paragraphs[0].alignment = 1
                
                # Add text to cells
                cell1.paragraphs[0].add_run(f"{school_name}")
                cell2.paragraphs[0].add_run(f"{exam_name}")
                
                # Add some space after the table
                doc.add_paragraph()
            
            doc.add_heading('H·ªç v√† T√™n: ...................................... L·ªõp: ...........', level=3)
            
            # Apply single line spacing with no spacing before/after
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing = 1.0  # Single spacing
            style.paragraph_format.space_before = 0
            style.paragraph_format.space_after = 0
            
            # Add page numbers to the footer
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = 2  # Right alignment
            run = paragraph.add_run()
            fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
            fldChar1.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)
            
            instrText = docx.oxml.shared.OxmlElement('w:instrText')
            instrText.set(docx.oxml.shared.qn('xml:space'), 'preserve')
            instrText.text = " PAGE "
            run._element.append(instrText)
            
            fldChar2 = docx.oxml.shared.OxmlElement('w:fldChar')
            fldChar2.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
            run._element.append(fldChar2)
            
            # Store answer key information for adding at the end
            answer_key = []
            
            # Add each question to the document
            for i, row in enumerate(st.session_state.randomized_questions.itertuples(), 1):
                # Create a single paragraph for question number and text
                question_para = doc.add_paragraph()
                question_para.paragraph_format.line_spacing = 1.0
                # Add space before if not the first question
                if i > 1:
                    question_para.paragraph_format.space_before = Pt(space_between_questions)
                else:
                    question_para.paragraph_format.space_before = 0
                question_para.paragraph_format.space_after = 0
                
                # Add the question number in bold
                question_number = question_para.add_run(f"C√¢u {i}: ")
                question_number.bold = True
                
                # Add the question text in italic
                question_text = question_para.add_run(row.Cau_hoi if hasattr(row, 'Cau_hoi') else row[2])
                question_text.italic = True
                
                # Get answer options and correct answer
                options = [
                    ("A", row.Tra_loi_1 if hasattr(row, 'Tra_loi_1') else row[3]),
                    ("B", row.Tra_loi_2 if hasattr(row, 'Tra_loi_2') else row[4]),
                    ("C", row.Tra_loi_3 if hasattr(row, 'Tra_loi_3') else row[5]),
                    ("D", row.Tra_loi_4 if hasattr(row, 'Tra_loi_4') else row[6])
                ]
                correct_answer = row.Dap_an_dung if hasattr(row, 'Dap_an_dung') else row[7]
                
                # Convert numeric correct answer to letter if needed
                if isinstance(correct_answer, (int, float)):
                    try:
                        numeric_answer = int(correct_answer)
                        if 1 <= numeric_answer <= 4:
                            correct_answer = chr(64 + numeric_answer)  # Convert 1->A, 2->B, 3->C, 4->D
                    except (ValueError, TypeError):
                        pass
                
                # Add to answer key
                answer_key.append(f"{i}{correct_answer}")
                
                # Calculate the layout based on answer lengths
                answer_lengths = [len(str(content)) for content in [opt[1] for opt in options]]
                avg_length = sum(answer_lengths) / len(answer_lengths)
                max_length = max(answer_lengths)
                
                # Decide layout based on answer lengths
                if max_length < 15 and avg_length < 10:  # Short answers - put all on one line
                    answers_para = doc.add_paragraph()
                    answers_para.paragraph_format.line_spacing = 1.0
                    answers_para.paragraph_format.space_before = 0
                    answers_para.paragraph_format.space_after = 0
                    
                    # Calculate spacing
                    page_width_points = 450
                    total_answer_text_len = sum(len(str(option_text)) for _, option_text in options)
                    total_letter_len = 8
                    space_to_distribute = max(20, page_width_points - total_answer_text_len - total_letter_len)
                    space_between = min(20, int(space_to_distribute / 3))
                    spacing = " " * space_between
                    
                    for idx, (option_letter, option_text) in enumerate(options):
                        if idx > 0:
                            answers_para.add_run(spacing)
                        
                        option_label = answers_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = answers_para.add_run(str(option_text))
                        if is_correct and underline_correct:
                            option_content.underline = True
                            
                elif max_length < 40 and avg_length < 30:  # Medium length - 2x2 layout
                    # First row (A and B)
                    row1_para = doc.add_paragraph()
                    row1_para.paragraph_format.line_spacing = 1.0
                    row1_para.paragraph_format.space_before = 0
                    row1_para.paragraph_format.space_after = 0
                    
                    page_width_points = 450
                    row1_text_len = sum(len(str(option_text)) for _, option_text in options[:2])
                    letter_len = 4
                    space_to_distribute = max(20, page_width_points - row1_text_len - letter_len)
                    space_between = min(20, int(space_to_distribute))
                    spacing = " " * space_between
                    
                    for idx in range(2):
                        if idx > 0:
                            row1_para.add_run(spacing)
                        
                        option_letter, option_text = options[idx]
                        option_label = row1_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = row1_para.add_run(str(option_text))
                        if is_correct and underline_correct:
                            option_content.underline = True
                    
                    # Second row (C and D)
                    row2_para = doc.add_paragraph()
                    row2_para.paragraph_format.line_spacing = 1.0
                    row2_para.paragraph_format.space_before = 0
                    row2_para.paragraph_format.space_after = 0
                    
                    row2_text_len = sum(len(str(option_text)) for _, option_text in options[2:4])
                    space_to_distribute = max(20, page_width_points - row2_text_len - letter_len)
                    space_between = min(20, int(space_to_distribute))
                    spacing = " " * space_between
                    
                    for idx in range(2, 4):
                        if idx > 2:
                            row2_para.add_run(spacing)
                        
                        option_letter, option_text = options[idx]
                        option_label = row2_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = row2_para.add_run(str(option_text))
                        if is_correct and underline_correct:
                            option_content.underline = True
                
                else:  # Long answers - one per line
                    for option_letter, option_text in options:
                        answer_para = doc.add_paragraph()
                        answer_para.paragraph_format.line_spacing = 1.0
                        answer_para.paragraph_format.space_before = 0
                        answer_para.paragraph_format.space_after = 0
                        
                        option_label = answer_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = answer_para.add_run(str(option_text))
                        if is_correct and underline_correct:
                            option_content.underline = True
            
            # Add the answer key at the end of the document
            if answer_key:
                doc.add_page_break()
                doc.add_heading('ƒê√°p √°n', level=1)
                
                answer_key_para = doc.add_paragraph()
                answer_key_para.paragraph_format.line_spacing = 1.0
                answer_key_para.paragraph_format.space_before = 0
                answer_key_para.paragraph_format.space_after = 0
                
                answers_per_line = 5
                page_width_points = 450
                entry_width = 3 * answers_per_line
                space_to_distribute = max(20, page_width_points - entry_width)
                space_between = min(20, int(space_to_distribute / (answers_per_line - 1)))
                spacing = " " * space_between
                
                for i, answer in enumerate(answer_key):
                    if i > 0 and i % answers_per_line == 0:
                        answer_key_para.add_run('\n')
                    elif i > 0:
                        answer_key_para.add_run(spacing)
                    answer_key_para.add_run(answer)
            
            # Save the document to a BytesIO object
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            
            # Convert the document to base64
            docx_base64 = base64.b64encode(docx_io.getvalue()).decode()
            
            # Create JavaScript to trigger download
            js = f"""
            <script>
                function downloadBase64File(base64Data, fileName) {{
                    const linkSource = `data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}`;
                    const downloadLink = document.createElement("a");
                    downloadLink.href = linkSource;
                    downloadLink.download = fileName + '.docx';
                    downloadLink.click();
                }}
                downloadBase64File('{docx_base64}', '{file_name}');
            </script>
            """
            
            st.success("T√†i li·ªáu Word ƒë√£ ƒë∆∞·ª£c t·∫°o th√†nh c√¥ng!")
            st.components.v1.html(js, height=0)

# Add instructions if no file has been uploaded
if not st.session_state.file_uploaded:
    st.info("Vui l√≤ng t·∫£i l√™n t·ªáp Excel ch·ª©a th∆∞ vi·ªán c√¢u h·ªèi ƒë·ªÉ b·∫Øt ƒë·∫ßu.")

# Add footer with instructions
st.markdown("---")
st.markdown("""
### H∆∞·ªõng d·∫´n s·ª≠ d·ª•ng:
1. T·∫£i l√™n t·ªáp Excel ch·ª©a th∆∞ vi·ªán c√¢u h·ªèi (ph·∫£i c√≥ √≠t nh·∫•t 10 c·ªôt).
2. Ch·ªçn s·ªë l∆∞·ª£ng c√¢u h·ªèi b·∫°n mu·ªën t·∫°o (m·∫∑c ƒë·ªãnh l√† 40).
3. X√°o tr·ªôn c√¢u h·ªèi v√† ƒë√°p √°n.
4a. ƒê·∫∑t kho·∫£ng c√°ch gi·ªØa c√°c c√¢u h·ªèi (0-30 ƒëi·ªÉm, trong ƒë√≥ 72 ƒëi·ªÉm = 1 inch).
4b. Ch·ªçn c√≥ g·∫°ch ch√¢n ƒë√°p √°n ƒë√∫ng hay kh√¥ng.
4c. Nh·∫≠p t√™n file Word.
4d. Nh·∫≠p t√™n tr∆∞·ªùng v√† t√™n b√†i ki·ªÉm tra.
5. Nh·∫•p v√†o "T·∫°o v√† T·∫£i xu·ªëng" ƒë·ªÉ t·∫°o t√†i li·ªáu Word.""")