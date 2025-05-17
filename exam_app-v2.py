import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt
import io
import tempfile
import os
import random
import base64

# Set page title and configuration
st.set_page_config(
    page_title="Phần Mềm Tạo Đề Thi",
    page_icon="📝",
    layout="wide"
)

# Title and description
st.title("Phần Mềm Tạo Đề Thi")
st.markdown("""
Ứng dụng này cho phép bạn:
1. Tải lên thư viện câu hỏi Excel của bạn
2. Chọn ngẫu nhiên một số câu hỏi
3. Xuất những câu hỏi đó ra tài liệu MS Word
""")

# Initialize session state variables
if 'question_library' not in st.session_state:
    st.session_state.question_library = None
if 'file_uploaded' not in st.session_state:
    st.session_state.file_uploaded = False
if 'file_valid' not in st.session_state:
    st.session_state.file_valid = False
if 'selected_questions' not in st.session_state:
    st.session_state.selected_questions = None
if 'randomized_questions' not in st.session_state:
    st.session_state.randomized_questions = None
if 'docx_data' not in st.session_state:
    st.session_state.docx_data = None
if 'trigger_download' not in st.session_state:
    st.session_state.trigger_download = False

# File upload section
st.header("Tải Lên Thư Viện Câu Hỏi")
st.markdown("Tải lên một tệp Excel (.xlsx hoặc .xls) chứa thư viện câu hỏi của bạn. Tệp phải có ít nhất 10 cột. Các cột theo thứ tự là: 'Ma cau', 'Cau hoi', 'Tra loi 1', 'Tra loi 2', 'Tra loi 3', 'Tra loi 4', 'Dap an dung', 'Bai', 'Phan', 'Do kho'.")

uploaded_file = st.file_uploader("Chọn tệp Excel", type=['xlsx', 'xls'])

if uploaded_file is not None:
    try:
        # Attempt to read the Excel file
        df = pd.read_excel(uploaded_file)
        
        # Define the column names
        column_names = [
            "Ma cau", "Cau hoi", "Tra loi 1", "Tra loi 2", "Tra loi 3", 
            "Tra loi 4", "Dap an dung", "Bai", "Phan", "Do kho"
        ]
        
        # Check if the file has at least 10 columns
        if len(df.columns) >= 10:
            # Select only the first 10 columns and rename them
            df = df.iloc[:, :10]
            df.columns = column_names
            st.success("Tệp đã được tải lên và kiểm tra thành công!")
            st.session_state.question_library = df
            st.session_state.file_uploaded = True
            st.session_state.file_valid = True
            
            # Display a preview of the uploaded data
            st.subheader("Xem trước thư viện câu hỏi")
            st.dataframe(df)
            
            # Display total number of questions in the library
            st.info(f"Tổng số câu hỏi trong thư viện: {len(df)}")
        else:
            st.error(f"Tệp đã tải lên có {len(df.columns)} cột. Nó phải có ít nhất 10 cột.")
            st.session_state.file_valid = False
    except Exception as e:
        if "Length mismatch" in str(e):
            st.error("Lỗi: Số lượng cột không khớp. Tệp Excel của bạn có quá nhiều cột. Vui lòng đảm bảo tệp có ít nhất 10 cột theo thứ tự: Ma cau, Cau hoi, Tra loi 1, Tra loi 2, Tra loi 3, Tra loi 4, Dap an dung, Bai, Phan, Do kho.")
        else:
            st.error(f"Lỗi khi đọc tệp: {str(e)}")
        st.session_state.file_valid = False

# Question Selection and Export section (only show if file is valid)
if st.session_state.file_valid:
    st.header("Tạo đề thi")
    
    # Step 1: Select random questions
    st.subheader("Bước 1: Chọn câu hỏi ngẫu nhiên")
    num_questions = st.number_input(
        "Số lượng câu hỏi cần chọn",
        min_value=1,
        max_value=len(st.session_state.question_library),
        value=min(40, len(st.session_state.question_library)),
        step=1
    )
    
    if st.button("Chọn câu hỏi ngẫu nhiên"):
        if num_questions > len(st.session_state.question_library):
            st.warning(f"Bạn đã yêu cầu {num_questions} câu hỏi, nhưng thư viện chỉ có {len(st.session_state.question_library)} câu hỏi. Tất cả câu hỏi sẽ được chọn.")
            selected_questions = st.session_state.question_library
        else:
            # Randomly select questions
            selected_indices = np.random.choice(
                len(st.session_state.question_library),
                size=num_questions,
                replace=False
            )
            selected_questions = st.session_state.question_library.iloc[selected_indices].reset_index(drop=True)
        
        # Store the selected questions in session state
        st.session_state.selected_questions = selected_questions
        st.session_state.randomized_questions = None  # Reset randomized questions
        st.session_state.docx_ready = None  # Reset docx ready state
        
        # Display the selected questions
        st.subheader("Các câu hỏi đã chọn")
        st.dataframe(selected_questions)
    
    # Step 2: Randomize questions and answers
    if st.session_state.selected_questions is not None:
        st.subheader("Bước 2: Xáo trộn câu hỏi và đáp án")
        if st.button("Xáo trộn câu hỏi và đáp án"):
            # Create a copy of selected questions
            randomized_questions = st.session_state.selected_questions.copy()
            
            # Randomize the order of questions
            randomized_questions = randomized_questions.sample(frac=1).reset_index(drop=True)
            
            # Randomize answers for each question
            for idx in range(len(randomized_questions)):
                # Get the original answers and correct answer
                original_answers = [
                    randomized_questions.iloc[idx]['Tra loi 1'],
                    randomized_questions.iloc[idx]['Tra loi 2'],
                    randomized_questions.iloc[idx]['Tra loi 3'],
                    randomized_questions.iloc[idx]['Tra loi 4']
                ]
                correct_answer = randomized_questions.iloc[idx]['Dap an dung']
                
                # Determine the original correct answer index
                correct_index = None
                if isinstance(correct_answer, str) and correct_answer.upper() in ["A", "B", "C", "D"]:
                    correct_index = ord(correct_answer.upper()) - ord('A')
                elif isinstance(correct_answer, (int, float)):
                    try:
                        temp_index = int(correct_answer) - 1
                        if 0 <= temp_index <= 3:
                            correct_index = temp_index
                    except (ValueError, TypeError):
                        pass
                
                if correct_index is None:
                    correct_index = 0
                
                # Get the correct answer content
                correct_content = original_answers[correct_index]
                
                # Shuffle the answers
                shuffled_answers = original_answers.copy()
                random.shuffle(shuffled_answers)
                
                # Find where the correct answer ended up after shuffling
                new_correct_index = shuffled_answers.index(correct_content)
                new_correct_letter = chr(65 + new_correct_index)
                
                # Update the question with shuffled answers and new correct answer
                randomized_questions.iloc[idx, 2:6] = shuffled_answers  # Update Tra loi 1-4
                randomized_questions.iloc[idx, 6] = new_correct_letter  # Update Dap an dung
            
            # Store randomized questions
            st.session_state.randomized_questions = randomized_questions
            
            # Display randomized questions
            st.subheader("Các câu hỏi đã xáo trộn")
            st.dataframe(randomized_questions)
    
    # Step 3: Export to Word
    if st.session_state.randomized_questions is not None:
        st.subheader("Bước 3: Xuất ra tài liệu Word")
        
        # Option to underline correct answers
        underline_correct = st.checkbox("Gạch chân đáp án đúng", value=True)
        
        # Input for space between questions
        space_between_questions = st.number_input(
            "Khoảng cách giữa các câu (điểm)", 
            min_value=0, 
            max_value=30,
            value=0,
            help="Đặt khoảng cách giữa các câu hỏi theo điểm (72 điểm = 1 inch). Giá trị cao hơn sẽ tạo ra khoảng cách lớn hơn giữa các câu hỏi."
        )
        
        if st.button("Tạo và tải xuống tài liệu Word"):
            # Create Word document
            doc = Document()
            doc.add_heading('Đề thi', level=1)
            
            # Apply single line spacing with no spacing before/after
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing = 1.0  # Single spacing
            style.paragraph_format.space_before = 0
            style.paragraph_format.space_after = 0
            
            # Store answer key information for adding at the end
            answer_key = []
            
            # Add each question to the document
            for i, row in enumerate(st.session_state.randomized_questions.itertuples(), 1):
                # Create a single paragraph for question number and text
                question_para = doc.add_paragraph()
                question_para.paragraph_format.line_spacing = 1.0
                # Add space before if not the first question
                if i > 1:
                    question_para.paragraph_format.space_before = Pt(space_between_questions)
                else:
                    question_para.paragraph_format.space_before = 0
                question_para.paragraph_format.space_after = 0
                
                # Add the question number in bold
                question_number = question_para.add_run(f"Câu {i}: ")
                question_number.bold = True
                
                # Add the question text in italic
                question_text = question_para.add_run(row.Cau_hoi if hasattr(row, 'Cau_hoi') else row[2])
                question_text.italic = True
                
                # Get answer options and correct answer
                options = [
                    ("A", row.Tra_loi_1 if hasattr(row, 'Tra_loi_1') else row[3]),
                    ("B", row.Tra_loi_2 if hasattr(row, 'Tra_loi_2') else row[4]),
                    ("C", row.Tra_loi_3 if hasattr(row, 'Tra_loi_3') else row[5]),
                    ("D", row.Tra_loi_4 if hasattr(row, 'Tra_loi_4') else row[6])
                ]
                correct_answer = row.Dap_an_dung if hasattr(row, 'Dap_an_dung') else row[7]
                
                # Add to answer key
                answer_key.append(f"{i}{correct_answer}")
                
                # Calculate the layout based on answer lengths
                answer_lengths = [len(content) for content in [opt[1] for opt in options]]
                avg_length = sum(answer_lengths) / len(answer_lengths)
                max_length = max(answer_lengths)
                
                # Decide layout based on answer lengths
                if max_length < 15 and avg_length < 10:  # Short answers - put all on one line
                    answers_para = doc.add_paragraph()
                    answers_para.paragraph_format.line_spacing = 1.0
                    answers_para.paragraph_format.space_before = 0
                    answers_para.paragraph_format.space_after = 0
                    
                    # Calculate spacing
                    page_width_points = 450
                    total_answer_text_len = sum(len(option_text) for _, option_text in options)
                    total_letter_len = 8
                    space_to_distribute = max(20, page_width_points - total_answer_text_len - total_letter_len)
                    space_between = min(20, int(space_to_distribute / 3))
                    spacing = " " * space_between
                    
                    for idx, (option_letter, option_text) in enumerate(options):
                        if idx > 0:
                            answers_para.add_run(spacing)
                        
                        option_label = answers_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = answers_para.add_run(option_text)
                        if is_correct and underline_correct:
                            option_content.underline = True
                            
                elif max_length < 40 and avg_length < 30:  # Medium length - 2x2 layout
                    # First row (A and B)
                    row1_para = doc.add_paragraph()
                    row1_para.paragraph_format.line_spacing = 1.0
                    row1_para.paragraph_format.space_before = 0
                    row1_para.paragraph_format.space_after = 0
                    
                    page_width_points = 450
                    row1_text_len = sum(len(option_text) for _, option_text in options[:2])
                    letter_len = 4
                    space_to_distribute = max(20, page_width_points - row1_text_len - letter_len)
                    space_between = min(20, int(space_to_distribute))
                    spacing = " " * space_between
                    
                    for idx in range(2):
                        if idx > 0:
                            row1_para.add_run(spacing)
                        
                        option_letter, option_text = options[idx]
                        option_label = row1_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = row1_para.add_run(option_text)
                        if is_correct and underline_correct:
                            option_content.underline = True
                    
                    # Second row (C and D)
                    row2_para = doc.add_paragraph()
                    row2_para.paragraph_format.line_spacing = 1.0
                    row2_para.paragraph_format.space_before = 0
                    row2_para.paragraph_format.space_after = 0
                    
                    row2_text_len = sum(len(option_text) for _, option_text in options[2:4])
                    space_to_distribute = max(20, page_width_points - row2_text_len - letter_len)
                    space_between = min(20, int(space_to_distribute))
                    spacing = " " * space_between
                    
                    for idx in range(2, 4):
                        if idx > 2:
                            row2_para.add_run(spacing)
                        
                        option_letter, option_text = options[idx]
                        option_label = row2_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = row2_para.add_run(option_text)
                        if is_correct and underline_correct:
                            option_content.underline = True
                
                else:  # Long answers - one per line
                    for option_letter, option_text in options:
                        answer_para = doc.add_paragraph()
                        answer_para.paragraph_format.line_spacing = 1.0
                        answer_para.paragraph_format.space_before = 0
                        answer_para.paragraph_format.space_after = 0
                        
                        option_label = answer_para.add_run(f"{option_letter}. ")
                        option_label.bold = True
                        
                        is_correct = (option_letter == correct_answer)
                        option_content = answer_para.add_run(option_text)
                        if is_correct and underline_correct:
                            option_content.underline = True
            
            # Add the answer key at the end of the document
            if answer_key:
                doc.add_page_break()
                doc.add_heading('Đáp án', level=1)
                
                answer_key_para = doc.add_paragraph()
                answer_key_para.paragraph_format.line_spacing = 1.0
                answer_key_para.paragraph_format.space_before = 0
                answer_key_para.paragraph_format.space_after = 0
                
                answers_per_line = 5
                page_width_points = 450
                entry_width = 3 * answers_per_line
                space_to_distribute = max(20, page_width_points - entry_width)
                space_between = min(20, int(space_to_distribute / (answers_per_line - 1)))
                spacing = " " * space_between
                
                for i, answer in enumerate(answer_key):
                    if i > 0 and i % answers_per_line == 0:
                        answer_key_para.add_run('\n')
                    elif i > 0:
                        answer_key_para.add_run(spacing)
                    answer_key_para.add_run(answer)
            
            # Save the document to a BytesIO object
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            
            # Convert the document to base64
            docx_base64 = base64.b64encode(docx_io.getvalue()).decode()
            
            # Create JavaScript to trigger download
            js = f"""
            <script>
                function downloadBase64File(base64Data, fileName) {{
                    const linkSource = `data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}`;
                    const downloadLink = document.createElement("a");
                    downloadLink.href = linkSource;
                    downloadLink.download = fileName;
                    downloadLink.click();
                }}
                downloadBase64File('{docx_base64}', 'examination_questions.docx');
            </script>
            """
            
            st.success("Tài liệu Word đã được tạo thành công!")
            st.components.v1.html(js, height=0)

# Add instructions if no file has been uploaded
if not st.session_state.file_uploaded:
    st.info("Vui lòng tải lên tệp Excel chứa thư viện câu hỏi để bắt đầu.")

# Add footer with instructions
st.markdown("---")
st.markdown("""
### Hướng dẫn sử dụng:
1. Tải lên tệp Excel chứa thư viện câu hỏi (phải có ít nhất 10 cột).
2. Chọn số lượng câu hỏi bạn muốn tạo (mặc định là 40).
3. Đặt khoảng cách giữa các câu hỏi (0-30 điểm, trong đó 72 điểm = 1 inch).
4. Chọn có gạch chân đáp án đúng hay không.
5. Nhấp vào "Tạo câu hỏi ngẫu nhiên" để chọn ngẫu nhiên câu hỏi từ thư viện.
6. Xem trước các câu hỏi đã chọn.
7. Tải xuống câu hỏi dưới dạng tài liệu Word.""")